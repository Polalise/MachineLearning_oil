import base64
import io
import json
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
ASSET_DIR = REPORT_DIR / "assets"
NOTEBOOK = ROOT / "model.ipynb"
DOCX_PATH = REPORT_DIR / "머신러닝 서비스모델 과제_유병현_초안.docx"
PDF_PATH = REPORT_DIR / "머신러닝 서비스모델 과제_유병현_초안.pdf"

TITLE = "Dubai 유가와 환율을 활용한 AI 기반\n소비자물가 상승률 예측모델 개발 및 시각화"
COURSE = "머신러닝 서비스모델 과제"
AUTHOR = "유병현"
REPORT_DATE = "2026년 6월 19일"


def extract_notebook_images():
    nb = json.loads(NOTEBOOK.read_text(encoding="utf-8"))
    rules = [
        ("fig, axes = plt.subplots(2, 2", "01_timeseries.png"),
        ("lag_correlations.plot", "02_lag_correlation.png"),
        ("sns.heatmap", "03_heatmap.png"),
        ("Persistence Baseline", "04_baseline.png"),
        ("residuals =", "05_prediction_residuals.png"),
        ("linear_coefficients =", "06_coefficients.png"),
        ("Model RMSE Comparison", "07_model_comparison.png"),
    ]
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    found = {}
    for cell in nb["cells"]:
        if cell.get("cell_type") != "code":
            continue
        source = "".join(cell.get("source", []))
        target_name = None
        for needle, filename in rules:
            if needle in source:
                target_name = filename
                break
        if not target_name:
            continue
        for output in cell.get("outputs", []):
            png = output.get("data", {}).get("image/png")
            if png:
                path = ASSET_DIR / target_name
                path.write_bytes(base64.b64decode(png))
                found[target_name] = path
                break
    return found


SCOPE_ROWS = [
    ["과제 구분", "내용"],
    ["데이터", "한국은행 ECOS 거시경제 지표 수집 및 월별 데이터셋 구축"],
    ["전처리", "가로형 자료의 세로형 변환, 환율 월평균, 결측치 점검, 파생변수 생성"],
    ["분석", "시계열 추이, 유가 시차별 상관관계, 변수 간 상관관계 분석"],
    ["모델링", "Persistence Baseline, Linear, Ridge, Random Forest, XGBoost 비교"],
    ["평가", "MAE, RMSE, R² 및 변수 제거 실험을 통한 추가 예측 기여 검증"],
    ["프로토타입", "Streamlit 기반 다음 달 소비자물가 상승률 예측 화면"],
]

DATA_ROWS = [
    ["데이터", "항목", "주기", "기간", "활용"],
    ["국제상품가격", "원유-Dubai", "월", "2005.01~2026.05", "유가 및 시차 변수"],
    ["소비자물가지수", "총지수(2020=100)", "월", "2005.01~2026.05", "전년동월비·타깃"],
    ["원/달러 환율", "매매기준율", "일→월", "2005.01~2026.06", "월평균·변동률"],
    ["기준금리", "한국은행 기준금리", "월", "2005.01~2026.05", "보조 거시변수"],
]

FEATURE_ROWS = [
    ["구분", "변수", "설명"],
    ["종속변수", "target_cpi_yoy", "다음 달 CPI 전년 동월 대비 상승률"],
    ["현재 물가", "cpi_yoy", "현재 CPI 전년 동월 대비 상승률"],
    ["유가", "oil_mom", "Dubai 유가 전월 대비 변동률"],
    ["유가 시차", "oil_lag_1/2/3/6", "유가 변동률의 1·2·3·6개월 시차"],
    ["환율", "fx_mom", "원/달러 월평균 환율의 전월 대비 변동률"],
    ["금리", "base_rate", "한국은행 기준금리"],
]

MODEL_ROWS = [
    ["Model", "MAE", "RMSE", "R²"],
    ["Linear Regression", "0.3036", "0.3602", "0.9278"],
    ["Ridge Regression", "0.3036", "0.3602", "0.9278"],
    ["Persistence Baseline", "0.3098", "0.3759", "0.9214"],
    ["Random Forest", "0.3374", "0.4248", "0.8996"],
    ["XGBoost", "0.4138", "0.5758", "0.8156"],
]

ABLATION_ROWS = [
    ["변수 구성", "MAE", "RMSE", "R²"],
    ["CPI Only", "0.3071", "0.3722", "0.9229"],
    ["CPI + FX + Rate", "0.3140", "0.3818", "0.9189"],
    ["Full Model (+ Oil)", "0.3036", "0.3602", "0.9278"],
]


def set_docx_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_docx_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_docx_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_docx_cell_text(
                cell,
                value,
                bold=r_idx == 0,
                color=(255, 255, 255) if r_idx == 0 else None,
            )
            if r_idx == 0:
                set_docx_cell_shading(cell, "1F4E78")
            elif r_idx % 2 == 0:
                set_docx_cell_shading(cell, "EAF2F8")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_docx_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_docx_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_docx_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_docx_image(doc, path, caption, width=16.5):
    if path and Path(path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width))
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(8)


def configure_docx(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9.5)
    for style_name, size, color in [
        ("Title", 25, (31, 78, 120)),
        ("Heading 1", 16, (31, 78, 120)),
        ("Heading 2", 12, (46, 116, 181)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)


def build_docx(images):
    doc = Document()
    configure_docx(doc)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(COURSE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(18)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(31, 78, 120)
    doc.add_paragraph()
    p = doc.add_paragraph("AI 개발 수행내역서")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(28)
    p.runs[0].bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(17)
    p.runs[0].bold = True
    for _ in range(3):
        doc.add_paragraph()
    add_docx_table(doc, [["담당자", AUTHOR], ["작성일", REPORT_DATE]], [3, 9])
    doc.add_page_break()

    add_docx_heading(doc, "AI 개발 수행내용", 1)
    add_docx_para(doc, f"사업과제: {TITLE.replace(chr(10), ' ')}")
    add_docx_heading(doc, "1. 개요 및 현황", 1)
    add_docx_heading(doc, "1.1 추진배경 및 목적", 2)
    add_docx_bullets(doc, [
        "한국은 원유의 대부분을 수입하므로 국제유가와 환율 변화가 국내 생산·운송비와 소비자물가에 영향을 줄 수 있다.",
        "원유가격 충격은 생산 및 유통 단계를 거쳐 소비자물가에 시차를 두고 반영될 가능성이 있다.",
        "월별 거시경제 데이터를 이용해 다음 달 소비자물가 상승률을 예측하고 유가 정보의 추가 예측 기여를 검증한다.",
        "예측 결과를 정책·기업 의사결정의 참고지표로 제시하고 Streamlit 기반 시범 서비스를 구현한다.",
    ])
    add_docx_heading(doc, "1.2 과제 범위", 2)
    add_docx_table(doc, SCOPE_ROWS, [3.2, 12.8])
    doc.add_page_break()

    add_docx_heading(doc, "1.3 과제 추진 방법", 2)
    add_docx_heading(doc, "1) 구축 대상 선정 기준", 2)
    add_docx_bullets(doc, [
        "데이터 접근성: 한국은행 ECOS에서 장기간의 공공 통계를 동일한 기준으로 확보할 수 있는가",
        "예측 가능성: 유가·환율·금리와 물가 사이의 시계열 관계를 월 단위로 구성할 수 있는가",
        "모델 효율성: 제한된 표본에서도 해석 가능한 회귀모델과 비교모델을 적용할 수 있는가",
        "활용성: 최신 데이터 갱신 시 다음 달 전망과 시나리오 분석으로 확장할 수 있는가",
    ])
    add_docx_heading(doc, "2) AI 예측 분석모델 적용 대상", 2)
    add_docx_table(doc, FEATURE_ROWS, [2.5, 4, 9.5])
    add_docx_heading(doc, "3) AI 분석모델 구축 프로세스", 2)
    add_docx_para(doc, "DATA IMPORTING → PREPROCESSING → EDA → TIME-SERIES SPLIT → MODEL COMPARISON → AUTO PREDICTION")
    add_docx_para(doc, "ECOS 자료를 월별로 통합하고 유가 시차 변수를 생성한 뒤, 과거 구간으로 학습하고 2022~2025년을 테스트하여 모델의 일반화 성능을 비교한다.")
    doc.add_page_break()

    add_docx_heading(doc, "연구개발 주요 결과물", 1)
    add_docx_heading(doc, "1. 데이터 수집", 1)
    add_docx_para(doc, "데이터 출처: 한국은행 경제통계시스템(ECOS), https://ecos.bok.or.kr")
    add_docx_table(doc, DATA_ROWS, [3, 3.2, 1.5, 3.2, 5.1])
    add_docx_heading(doc, "2. 데이터 전처리", 1)
    add_docx_bullets(doc, [
        "가로형 CSV에서 원자료 행만 선택하고 날짜 열을 세로형으로 변환",
        "일별 원/달러 환율의 천 단위 쉼표 제거 후 월평균 계산",
        "네 데이터셋을 날짜 기준으로 병합하여 257개월, 5개 기본 변수 구성",
        "CPI 전년동월비, 유가·환율 전월비, 유가 1·2·3·6개월 시차 생성",
        "시차 계산에서 발생한 초기 결측치만 모델링 단계에서 제거",
    ])
    add_docx_para(doc, "최종 모델링 표본은 239개월이며 학습 191개월, 테스트 48개월로 시간 순서에 따라 분리하였다.")
    doc.add_page_break()

    add_docx_heading(doc, "3. 데이터 분석", 1)
    add_docx_heading(doc, "3.1 기술통계 및 시계열 추이", 2)
    add_docx_para(doc, "2008년 금융위기, 2020년 코로나19, 2022년 원자재가격 상승기 등 실제 경제 충격을 포함한다. 이러한 극단 구간은 오류가 아닌 경제 현상이므로 제거하지 않았다.")
    add_docx_image(doc, images.get("01_timeseries.png"), "[그림 1] Dubai 유가, CPI 상승률, 환율 및 기준금리 추이")
    doc.add_page_break()

    add_docx_heading(doc, "3.2 유가 시차별 상관관계", 2)
    add_docx_para(doc, "6개월 시차 유가 변동률과 다음 달 물가상승률의 상관계수가 약 0.19로 가장 높았다. 이는 물가 반영에 시차가 존재할 가능성을 보여주지만 상관관계만으로 인과성을 단정할 수는 없다.")
    add_docx_image(doc, images.get("02_lag_correlation.png"), "[그림 2] 유가 시차별 다음 달 물가상승률 상관관계", 15)
    add_docx_heading(doc, "3.3 전체 변수 상관관계", 2)
    add_docx_para(doc, "현재 물가상승률과 다음 달 상승률의 상관계수는 0.96으로 물가의 강한 자기 지속성을 보였다. 유가 시차 변수 간 상관계수는 최대 약 0.33으로 심각한 다중공선성은 확인되지 않았다.")
    add_docx_image(doc, images.get("03_heatmap.png"), "[그림 3] 모델 입력변수 상관관계 히트맵", 14)
    doc.add_page_break()

    add_docx_heading(doc, "4. 데이터 학습 및 모델 정의", 1)
    add_docx_heading(doc, "4.1 학습·테스트 분리", 2)
    add_docx_para(doc, "무작위 분할로 인한 미래정보 누수를 방지하기 위해 예측 대상 월을 기준으로 학습(2006.02~2021.12)과 테스트(2022.01~2025.12)를 구분하였다.")
    add_docx_heading(doc, "4.2 기준 모델 및 후보 모델", 2)
    add_docx_bullets(doc, [
        "Persistence Baseline: 현재 물가상승률이 다음 달에도 유지된다고 가정",
        "Linear Regression: 표준화된 입력변수의 선형 관계와 해석력 활용",
        "Ridge Regression: 시계열 교차검증으로 규제 강도 최적화",
        "Random Forest / XGBoost: 비선형 관계와 변수 상호작용 비교",
    ])
    add_docx_image(doc, images.get("04_baseline.png"), "[그림 4] 실제값과 지속성 기준 모델 비교")
    doc.add_page_break()

    add_docx_heading(doc, "4.3 모델 성능평가 및 최종 선정", 2)
    add_docx_table(doc, MODEL_ROWS, [5.8, 3.2, 3.2, 3.2])
    add_docx_para(doc, "Linear Regression은 RMSE 0.3602, R² 0.9278로 가장 우수했다. Ridge의 최적 alpha는 0.001로 일반 선형회귀와 사실상 동일했으므로 단순성과 해석 가능성을 고려해 Linear Regression을 최종 모델로 선정하였다.")
    add_docx_image(doc, images.get("07_model_comparison.png"), "[그림 5] 모델별 RMSE 및 R² 비교")
    add_docx_heading(doc, "4.4 유가 변수 추가 기여 검증", 2)
    add_docx_table(doc, ABLATION_ROWS, [7, 3, 3, 3])
    add_docx_para(doc, "유가 변수를 포함한 Full Model은 유가 제외 모델보다 RMSE를 약 5.6% 줄였다. 따라서 유가 정보가 현재 CPI의 자기 지속성 외에 소폭의 추가 예측 정보를 제공한 것으로 판단하였다.")
    doc.add_page_break()

    add_docx_heading(doc, "5. 최종 모델 예측 및 해석", 1)
    add_docx_heading(doc, "5.1 실제값과 예측값", 2)
    add_docx_para(doc, "최종 모델은 전반적인 물가 흐름을 추종했으나 급등 초기에는 상승 폭을 과소 예측하고 전환점에서는 실제 변화보다 늦게 반응하는 경향을 보였다.")
    add_docx_image(doc, images.get("05_prediction_residuals.png"), "[그림 6] Linear Regression 예측 및 잔차")
    add_docx_heading(doc, "5.2 표준화 회귀계수", 2)
    add_docx_para(doc, "현재 물가상승률의 계수가 가장 컸으며 유가 변수 중에서는 당월 변동률과 6개월 시차 변동률의 기여가 상대적으로 높았다. 계수는 예측 관계이며 직접적인 인과효과가 아니다.")
    add_docx_image(doc, images.get("06_coefficients.png"), "[그림 7] 표준화 선형회귀 계수", 14.5)
    doc.add_page_break()

    add_docx_heading(doc, "6. 프로토타이핑(화면)", 1)
    add_docx_heading(doc, "6.1 Streamlit 기반 다음 달 물가 예측", 2)
    add_docx_bullets(doc, [
        "최신 CPI 상승률, Dubai 유가 변동률, 환율 변동률 및 기준금리 입력",
        "유가 1·2·3·6개월 시차는 최신 관측값을 자동 적용",
        "예측 대상 월과 예상 CPI 상승률, 최종 모델 R² 표시",
        "소비자물가 상승률 및 Dubai 유가의 장기 추이 제공",
    ])
    add_docx_image(doc, ASSET_DIR / "streamlit_full.png", "[그림 8] 소비자물가 상승률 예측 프로토타입", 16.5)
    add_docx_para(doc, "2026년 5월까지의 최신 정보를 입력한 결과 2026년 6월 CPI 전년 동월 대비 상승률은 3.28%로 예측되었다. 이는 미래 예측치이며 테스트 성능평가 결과와 구분한다.")
    doc.add_page_break()

    add_docx_heading(doc, "7. 결론 및 향후 계획", 1)
    add_docx_heading(doc, "7.1 모델 개발 결과", 2)
    add_docx_bullets(doc, [
        "현재 CPI 상승률이 다음 달 물가를 예측하는 가장 중요한 변수로 확인되었다.",
        "Dubai 유가의 당월 변동률과 6개월 시차 변동률이 보조적인 예측 정보를 제공하였다.",
        "Linear Regression이 MAE 0.3036, RMSE 0.3602, R² 0.9278로 가장 우수했다.",
        "유가 변수를 추가했을 때 유가 제외 모델보다 RMSE가 약 5.6% 개선되었다.",
    ])
    add_docx_heading(doc, "7.2 한계", 2)
    add_docx_bullets(doc, [
        "월별 표본이 제한적이며 금융위기·감염병·지정학적 충격 등 구조 변화가 포함되어 있다.",
        "임금, 수요, 공급망, 재정정책, 기대인플레이션 등 다른 물가 요인을 포함하지 못했다.",
        "현재 CPI의 자기 지속성이 예측력의 대부분을 차지하며 유가·환율의 개선 폭은 크지 않다.",
        "상관관계와 회귀계수는 유가의 직접적인 인과효과를 의미하지 않는다.",
    ])
    add_docx_heading(doc, "7.3 향후 개선 방향", 2)
    add_docx_bullets(doc, [
        "수입물가지수, 생산자물가지수, 임금 및 기대인플레이션 변수 추가",
        "원화 환산 유가 변동률과 유가·환율 상호작용 검토",
        "롤링 윈도우 재학습 및 구조 변화 탐지",
        "ECOS API 기반 월별 자동 수집·재학습 파이프라인 구축",
    ])
    add_docx_heading(doc, "참고자료", 2)
    add_docx_para(doc, "한국은행 경제통계시스템(ECOS), 소비자물가지수·국제상품가격·환율·기준금리 통계")
    add_docx_para(doc, "U.S. Energy Information Administration, South Korea Country Analysis Brief")
    add_docx_para(doc, "GS칼텍스 미디어허브, 국제유가와 중동·아시아 원유시장 분석")

    doc.save(DOCX_PATH)


def pdf_styles():
    font_regular = Path(r"C:\Windows\Fonts\malgun.ttf")
    font_bold = Path(r"C:\Windows\Fonts\malgunbd.ttf")
    pdfmetrics.registerFont(TTFont("Malgun", str(font_regular)))
    pdfmetrics.registerFont(TTFont("MalgunB", str(font_bold)))
    return {
        "title": ParagraphStyle("TitleK", fontName="MalgunB", fontSize=23, leading=32, alignment=TA_CENTER, textColor=colors.HexColor("#1F4E78")),
        "cover": ParagraphStyle("CoverK", fontName="MalgunB", fontSize=17, leading=24, alignment=TA_CENTER),
        "h1": ParagraphStyle("H1K", fontName="MalgunB", fontSize=15, leading=20, spaceBefore=7, spaceAfter=6, textColor=colors.HexColor("#1F4E78")),
        "h2": ParagraphStyle("H2K", fontName="MalgunB", fontSize=11.5, leading=16, spaceBefore=6, spaceAfter=4, textColor=colors.HexColor("#2E75B5")),
        "body": ParagraphStyle("BodyK", fontName="Malgun", fontSize=8.8, leading=13, spaceAfter=4),
        "small": ParagraphStyle("SmallK", fontName="Malgun", fontSize=7.7, leading=10, alignment=TA_CENTER, textColor=colors.HexColor("#555555")),
        "bullet": ParagraphStyle("BulletK", fontName="Malgun", fontSize=8.7, leading=12.5, leftIndent=10, firstLineIndent=-6, bulletIndent=2, spaceAfter=2),
        "table": ParagraphStyle("TableK", fontName="Malgun", fontSize=7.3, leading=9),
        "table_h": ParagraphStyle("TableHK", fontName="MalgunB", fontSize=7.3, leading=9, alignment=TA_CENTER, textColor=colors.white),
    }


def page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Malgun", 7.5)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(18 * mm, 10 * mm, COURSE)
    canvas.drawRightString(192 * mm, 10 * mm, str(doc.page))
    canvas.restoreState()


def pdf_table(rows, styles, widths):
    data = []
    for r_idx, row in enumerate(rows):
        style = styles["table_h"] if r_idx == 0 else styles["table"]
        data.append([Paragraph(str(v), style) for v in row])
    table = Table(data, colWidths=[w * mm for w in widths], repeatRows=1, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#AAB7C4")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAF2F8")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def pdf_image(path, caption, styles, max_w=172, max_h=122):
    items = []
    if path and Path(path).exists():
        with PILImage.open(path) as im:
            w, h = im.size
        scale = min(max_w * mm / w, max_h * mm / h)
        items.append(Image(str(path), width=w * scale, height=h * scale, hAlign="CENTER"))
        items.append(Spacer(1, 2 * mm))
        items.append(Paragraph(caption, styles["small"]))
    return items


def build_pdf(images):
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_PATH), pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=17 * mm, bottomMargin=16 * mm,
        title=COURSE, author=AUTHOR,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates(PageTemplate(id="all", frames=frame, onPage=page_number))
    story = []
    H1 = lambda t: story.append(Paragraph(t, styles["h1"]))
    H2 = lambda t: story.append(Paragraph(t, styles["h2"]))
    P = lambda t: story.append(Paragraph(t, styles["body"]))
    B = lambda items: [story.append(Paragraph("• " + x, styles["bullet"])) for x in items]

    story += [Spacer(1, 34 * mm), Paragraph(COURSE, styles["cover"]), Spacer(1, 12 * mm), Paragraph("AI 개발 수행내역서", styles["title"]), Spacer(1, 18 * mm), Paragraph(TITLE.replace("\n", "<br/>"), styles["cover"]), Spacer(1, 30 * mm)]
    story.append(pdf_table([["담당자", AUTHOR], ["작성일", REPORT_DATE]], styles, [35, 100]))
    story.append(PageBreak())

    H1("AI 개발 수행내용")
    P("<b>사업과제:</b> " + TITLE.replace("\n", " "))
    H1("1. 개요 및 현황")
    H2("1.1 추진배경 및 목적")
    B([
        "한국은 원유의 대부분을 수입하므로 국제유가와 환율 변화가 국내 생산·운송비와 소비자물가에 영향을 줄 수 있다.",
        "원유가격 충격은 생산 및 유통 단계를 거쳐 소비자물가에 시차를 두고 반영될 가능성이 있다.",
        "월별 거시경제 데이터로 다음 달 소비자물가 상승률을 예측하고 유가 정보의 추가 예측 기여를 검증한다.",
        "Streamlit 기반 시범 서비스를 통해 최신 전망과 시나리오 분석 기능을 제공한다.",
    ])
    H2("1.2 과제 범위")
    story.append(pdf_table(SCOPE_ROWS, styles, [31, 142]))
    story.append(PageBreak())

    H2("1.3 과제 추진 방법")
    H2("1) 구축 대상 선정 기준")
    B([
        "데이터 접근성: 한국은행 ECOS의 장기 공공 통계를 동일 기준으로 확보",
        "예측 가능성: 유가·환율·금리와 물가의 월별 시계열 관계 구성",
        "모델 효율성: 제한된 표본에서 해석 가능한 회귀모델과 비교모델 적용",
        "활용성: 최신 자료 갱신 시 다음 달 전망과 시나리오 분석으로 확장",
    ])
    H2("2) AI 예측 분석모델 적용 대상")
    story.append(pdf_table(FEATURE_ROWS, styles, [25, 39, 109]))
    H2("3) AI 분석모델 구축 프로세스")
    P("<b>DATA IMPORTING → PREPROCESSING → EDA → TIME-SERIES SPLIT → MODEL COMPARISON → AUTO PREDICTION</b>")
    P("ECOS 자료를 월별로 통합하고 시차 변수를 생성한 뒤, 과거 구간으로 학습하고 2022~2025년을 테스트하여 일반화 성능을 비교한다.")
    story.append(PageBreak())

    H1("연구개발 주요 결과물")
    H1("1. 데이터 수집")
    P("데이터 출처: 한국은행 경제통계시스템(ECOS), https://ecos.bok.or.kr")
    story.append(pdf_table(DATA_ROWS, styles, [30, 34, 17, 37, 55]))
    H1("2. 데이터 전처리")
    B([
        "가로형 CSV에서 원자료 행을 선택하고 날짜 열을 세로형으로 변환",
        "일별 환율의 천 단위 쉼표 제거 후 월평균 계산",
        "네 자료를 날짜 기준으로 병합하여 257개월의 기본 데이터셋 구성",
        "CPI 전년동월비, 유가·환율 전월비, 유가 1·2·3·6개월 시차 생성",
        "최종 239개월을 학습 191개월, 테스트 48개월로 시간 순서에 따라 분리",
    ])
    story.append(PageBreak())

    H1("3. 데이터 분석")
    H2("3.1 기술통계 및 시계열 추이")
    P("2008년 금융위기, 2020년 코로나19, 2022년 원자재가격 상승기 등 실제 경제 충격을 포함하며 극단 구간을 임의로 제거하지 않았다.")
    story += pdf_image(images.get("01_timeseries.png"), "[그림 1] 주요 경제지표의 시계열 추이", styles, 172, 125)
    story.append(PageBreak())

    H2("3.2 유가 시차별 상관관계")
    P("6개월 시차 유가 변동률이 다음 달 물가상승률과 약 0.19의 가장 높은 상관관계를 보였다. 이는 시차 가능성을 시사하지만 인과성을 증명하지는 않는다.")
    story += pdf_image(images.get("02_lag_correlation.png"), "[그림 2] 유가 시차별 상관관계", styles, 158, 82)
    H2("3.3 전체 변수 상관관계")
    P("현재 물가상승률과 다음 달 상승률은 0.96의 강한 자기상관을 보였고, 유가 시차 변수 간 심각한 다중공선성은 확인되지 않았다.")
    story += pdf_image(images.get("03_heatmap.png"), "[그림 3] 입력변수 상관관계 히트맵", styles, 135, 105)
    story.append(PageBreak())

    H1("4. 데이터 학습 및 모델 정의")
    H2("4.1 학습·테스트 분리")
    P("미래정보 누수를 막기 위해 학습 2006.02~2021.12, 테스트 2022.01~2025.12로 분리하였다.")
    H2("4.2 기준 모델 및 후보 모델")
    B([
        "Persistence Baseline: 현재 물가상승률을 다음 달 예측값으로 사용",
        "Linear/Ridge: 표준화된 변수의 선형 관계와 해석력 활용",
        "Random Forest/XGBoost: 비선형 관계 및 상호작용 비교",
    ])
    story += pdf_image(images.get("04_baseline.png"), "[그림 4] 실제값과 지속성 기준 모델", styles, 172, 72)
    story.append(PageBreak())

    H2("4.3 모델 성능평가 및 최종 선정")
    story.append(pdf_table(MODEL_ROWS, styles, [70, 34, 34, 34]))
    P("Linear Regression이 RMSE 0.3602, R² 0.9278로 가장 우수했다. Ridge는 사실상 같은 성능이었으나 규제의 이점이 없어 Linear Regression을 최종 선정하였다.")
    story += pdf_image(images.get("07_model_comparison.png"), "[그림 5] 모델별 RMSE와 R² 비교", styles, 172, 65)
    H2("4.4 유가 변수 추가 기여")
    story.append(pdf_table(ABLATION_ROWS, styles, [70, 34, 34, 34]))
    P("Full Model은 유가 제외 모델보다 RMSE를 약 5.6% 줄여 유가 정보의 소폭 추가 예측 기여를 확인하였다.")
    story.append(PageBreak())

    H1("5. 최종 모델 예측 및 해석")
    H2("5.1 실제값·예측값 및 잔차")
    P("전체 추세는 잘 추종했으나 급등 초기에는 상승 폭을 과소 예측하고 전환점에서 늦게 반응하는 경향을 보였다.")
    story += pdf_image(images.get("05_prediction_residuals.png"), "[그림 6] 최종 모델 예측과 잔차", styles, 172, 110)
    H2("5.2 표준화 회귀계수")
    P("현재 CPI 상승률이 가장 큰 계수를 보였으며 유가 당월 변동률과 6개월 시차 변동률이 보조적으로 기여했다.")
    story += pdf_image(images.get("06_coefficients.png"), "[그림 7] 표준화 회귀계수", styles, 145, 82)
    story.append(PageBreak())

    H1("6. 프로토타이핑(화면)")
    H2("6.1 Streamlit 기반 다음 달 물가 예측")
    B([
        "최신 CPI, Dubai 유가, 환율 변동률 및 기준금리 입력",
        "유가 시차는 최신 관측값을 자동 적용",
        "예측 대상 월, 예상 CPI 상승률, 모델 R² 표시",
        "CPI 상승률과 Dubai 유가 장기 추이 제공",
    ])
    story += pdf_image(ASSET_DIR / "streamlit_full.png", "[그림 8] Streamlit 소비자물가 예측 화면", styles, 172, 120)
    P("최신 자료 기준 2026년 6월 CPI 상승률은 3.28%로 예측되었다. 이는 미래 예측치이며 테스트 평가와 구분한다.")
    story.append(PageBreak())

    H1("7. 결론 및 향후 계획")
    H2("7.1 모델 개발 결과")
    B([
        "현재 CPI 상승률이 다음 달 물가의 가장 중요한 예측변수였다.",
        "유가 당월 변동률과 6개월 시차 변동률이 추가 정보를 제공했다.",
        "Linear Regression은 MAE 0.3036, RMSE 0.3602, R² 0.9278을 기록했다.",
        "유가 추가 시 유가 제외 모델보다 RMSE가 약 5.6% 개선되었다.",
    ])
    H2("7.2 한계")
    B([
        "월별 표본이 제한적이며 금융위기·감염병·지정학적 충격을 포함한다.",
        "임금, 수요, 공급망, 기대인플레이션 등 다른 물가 요인을 포함하지 못했다.",
        "예측력 대부분은 현재 CPI의 자기 지속성에서 나온다.",
        "상관관계와 회귀계수는 직접적인 인과효과를 의미하지 않는다.",
    ])
    H2("7.3 향후 개선 방향")
    B([
        "수입·생산자물가, 임금, 기대인플레이션 변수 추가",
        "원화 환산 유가와 유가·환율 상호작용 검토",
        "롤링 윈도우 재학습 및 구조 변화 탐지",
        "ECOS API 기반 자동 수집·재학습 구축",
    ])
    H2("참고자료")
    P("한국은행 ECOS; U.S. EIA South Korea Country Analysis Brief; GS칼텍스 미디어허브 국제유가 분석")

    doc.build(story)


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    images = extract_notebook_images()
    build_docx(images)
    build_pdf(images)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
