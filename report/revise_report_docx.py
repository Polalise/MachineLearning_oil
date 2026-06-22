from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).parent
SOURCE = ROOT / "머신러닝 서비스모델 과제_유병현_초안.docx"
OUTPUT = ROOT / "머신러닝 서비스모델 과제_유병현_양식반영본.docx"


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def replace_paragraph(paragraph, text, style=None):
    paragraph.clear()
    paragraph.add_run(text)
    if style:
        paragraph.style = style


def insert_after(paragraph, text="", style=None):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def add_material_marker(paragraph, description):
    marker = insert_after(paragraph)
    run = marker.add_run(f"[사용자 자료 삽입 필요] {description}")
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return marker


document = Document(SOURCE)

# Align the opening section numbering with the supplied report template.
project_paragraph = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.strip().startswith("사업과제:")
)
project_text = project_paragraph.text.strip()
replace_paragraph(project_paragraph, "1. 사업과제", "Heading 1")
insert_after(project_paragraph, project_text.removeprefix("사업과제:").strip())

heading_updates = {
    "1. 개요 및 현황": "2. 개요 및 현황",
    "1.1 추진배경 및 목적": "2.1 추진배경 및 목적",
    "1.2 과제 범위": "2.2 과제 범위",
    "1.3 과제 추진 방법": "2.3 과제 추진 방법",
}
for old_text, new_text in heading_updates.items():
    replace_paragraph(find_paragraph(document, old_text), new_text)

# Keep the report concise while making missing evidence explicit.
source_note = find_paragraph(
    document,
    "데이터 출처: 한국은행 경제통계시스템(ECOS), https://ecos.bok.or.kr",
)
add_material_marker(
    source_note,
    "한국은행 ECOS 데이터 조회 화면 또는 다운로드한 원자료 일부 캡처",
)

preprocess_summary = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.strip().startswith("최종 모델링 표본은 239개월")
)
add_material_marker(
    preprocess_summary,
    "model.ipynb의 통합 데이터 크기·결측치 확인 결과 화면",
)

final_model_summary = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.strip().startswith("Linear Regression은 RMSE 0.3602")
)
add_material_marker(
    final_model_summary,
    "최종 Linear Regression 모델 정의·학습 코드 실행 화면",
)

# Preserve the existing header/footer and section settings explicitly.
for source_section, output_section in zip(Document(SOURCE).sections, document.sections):
    output_section.header_distance = source_section.header_distance
    output_section.footer_distance = source_section.footer_distance
    for source_paragraph, output_paragraph in zip(
        source_section.header.paragraphs,
        output_section.header.paragraphs,
    ):
        output_paragraph._p.getparent().replace(
            output_paragraph._p,
            deepcopy(source_paragraph._p),
        )

document.save(OUTPUT)
print(OUTPUT)
